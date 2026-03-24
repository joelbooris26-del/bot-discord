from flask import Flask
from threading import Thread

app = Flask('')

@app.route('/')
def home():
    return "Bot activo"

def run():
    app.run(host='0.0.0.0', port=8080)

def keep_alive():
    t = Thread(target=run)
    t.start()

import discord
from discord.ext import commands
from docx import Document
import os
import shutil
import json
from datetime import datetime

# ---------------- CONFIG ----------------
TOKEN = os.environ.get("DISCORD_TOKEN")

OWNER_ID = 583251995729723393
LOG_CHANNEL_ID = 1485699759736885318
LOG_BORRADOS_ID = 1485706506207756499

DOC_FILE = "registro_clientes.docx"
DATA_FILE = "data.json"
QUEUE_FILE = "cola.txt"
QUEUE_MESSAGE_ID_FILE = "queue_msg_id.txt"

intents = discord.Intents.default()
intents.message_content = True
intents.members = True

bot = commands.Bot(command_prefix="!", intents=intents)

@bot.event
async def on_ready():
    print(f"✅ Bot conectado como {bot.user}")

# ---------------- FUNCIONES ----------------
def es_owner(ctx):
    return ctx.author.id == OWNER_ID


def backup_doc():
    if os.path.exists(DOC_FILE):
        shutil.copy(DOC_FILE, "backup_registro.docx")


def crear_doc_si_no_existe():
    if not os.path.exists(DOC_FILE):
        doc = Document()
        doc.add_heading('Registro de Clientes', 0)
        table = doc.add_table(rows=1, cols=5)
        headers = ["Usuario", "ID", "Producto", "Precio", "Fecha"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        doc.save(DOC_FILE)


def cargar_datos():
    if not os.path.exists(DATA_FILE):
        return []
    with open(DATA_FILE, "r", encoding="utf-8") as f:
        return json.load(f)


def guardar_datos(data):
    with open(DATA_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=4)


def leer_cola():
    if not os.path.exists(QUEUE_FILE):
        return []
    with open(QUEUE_FILE, "r", encoding="utf-8") as f:
        return f.readlines()


def guardar_cola(lineas):
    with open(QUEUE_FILE, "w", encoding="utf-8") as f:
        f.writelines(lineas)


async def actualizar_mensaje_cola(channel):
    cola = leer_cola()
    texto = "\n".join([f"🔹 **{i+1}.** {line.strip()}" for i, line in enumerate(cola)]) or "Sin pedidos"

    embed = discord.Embed(
        title="📋 Cola de pedidos",
        description=texto,
        color=discord.Color.orange()
    )

    try:
        # intentar leer ID
        if os.path.exists(QUEUE_MESSAGE_ID_FILE):
            with open(QUEUE_MESSAGE_ID_FILE, "r") as f:
                msg_id = int(f.read())

            msg = await channel.fetch_message(msg_id)
            await msg.edit(embed=embed)
            return
    except:
        pass

    # si falla → crear nuevo
    msg = await channel.send(embed=embed)

    with open(QUEUE_MESSAGE_ID_FILE, "w") as f:
        f.write(str(msg.id))

# ---------------- REGISTRAR ----------------
@bot.command()
async def registrar(ctx, usuario: str, user_id: str, precio: str, *, producto: str):
    if not es_owner(ctx):
        return await ctx.send("❌ No tienes permiso")

    crear_doc_si_no_existe()

    try:
        precio_val = float(precio.replace("€", "").replace(",", "."))
    except:
        return await ctx.send("❌ Precio inválido")

    doc = Document(DOC_FILE)
    table = doc.tables[0]

    row = table.add_row().cells
    row[0].text = usuario
    row[1].text = user_id
    row[2].text = producto
    row[3].text = f"{precio_val}€"
    row[4].text = datetime.now().strftime("%d/%m/%Y")

    doc.save(DOC_FILE)
    backup_doc()

    data = cargar_datos()
    data.append({
        "usuario": usuario,
        "id": user_id,
        "producto": producto,
        "precio": precio_val,
        "fecha": datetime.now().strftime("%d/%m/%Y")
    })
    guardar_datos(data)

    embed = discord.Embed(
        title="✅ Compra registrada",
        color=discord.Color.green()
    )
    embed.add_field(name="👤 Usuario", value=usuario, inline=False)
    embed.add_field(name="📦 Producto", value=producto, inline=False)
    embed.add_field(name="💰 Precio", value=f"{precio_val}€", inline=False)

    await ctx.send(embed=embed)

    log_channel = bot.get_channel(LOG_CHANNEL_ID)
    if log_channel and hasattr(log_channel, "send"):
        await log_channel.send(embed=embed)

# ---------------- STATS ----------------
@bot.command()
async def stats(ctx):
    if not es_owner(ctx):
        return await ctx.send("❌ No tienes permiso")

    data = cargar_datos()

    if not data:
        return await ctx.send("❌ No hay datos")

    total_pedidos = len(data)
    total_dinero = sum(d["precio"] for d in data)

    embed = discord.Embed(
        title="📊 Estadísticas",
        color=discord.Color.blue()
    )
    embed.add_field(name="📦 Pedidos", value=total_pedidos, inline=True)
    embed.add_field(name="💰 Total", value=f"{total_dinero}€", inline=True)

    await ctx.send(embed=embed)

# ---------------- COLA ----------------

@bot.command()
async def setup(ctx):
    if not es_owner(ctx):
        return await ctx.send("❌ No tienes permiso")

    await actualizar_mensaje_cola(ctx.channel)

    await ctx.send("✅ Cola configurada correctamente", delete_after=5)
    
    
@bot.command()
async def añadir_cola(ctx, user: discord.Member, *, producto: str):
    if not es_owner(ctx):
        return

    await ctx.message.delete()  # 👈 AÑADE ESTO

    linea = f"{user.name} | ID: {user.id} | Producto: {producto}\n"

    with open(QUEUE_FILE, "a", encoding="utf-8") as f:
        f.write(linea)

    cola = leer_cola()
    posicion = len(cola)

    await actualizar_mensaje_cola(ctx.channel)

    embed = discord.Embed(
        title="✅ Añadido a la cola",
        description=f"{user.name} está en el puesto #{posicion}",
        color=discord.Color.green()
    )

    await ctx.send(embed=embed, delete_after=5)


@bot.command()
async def borrar_cola(ctx, posicion: int):
    if not es_owner(ctx):
        return

    cola = leer_cola()
    if posicion < 1 or posicion > len(cola):
        return await ctx.send("❌ Posición inválida")

    eliminado = cola.pop(posicion - 1)
    guardar_cola(cola)

    await actualizar_mensaje_cola(ctx.channel)

    nombre = eliminado.split(" | ")[0]

    embed = discord.Embed(
        title="🗑️ Pedido eliminado",
        description=f"{nombre} eliminado de la cola",
        color=discord.Color.red()
    )

    await ctx.send(embed=embed)

# ---------------- BORRAR REGISTRO ----------------
@bot.command()
async def borrar_registro(ctx, user_id: str, *, filtro: str = None):
    if not es_owner(ctx):
        return await ctx.send("❌ No tienes permiso")

    await ctx.send("⚠️ Escribe 'confirmar' para continuar")

    def check(m):
        return m.author == ctx.author and m.channel == ctx.channel

    try:
        msg = await bot.wait_for("message", check=check, timeout=15)
        if msg.content.lower() != "confirmar":
            return await ctx.send("❌ Cancelado")
    except:
        return await ctx.send("⏰ Tiempo agotado")

    data = cargar_datos()
    nuevos = []
    eliminados = []

    for d in data:
        if d["id"] == user_id:
            if filtro == "todos" or (filtro and filtro.lower() in d["producto"].lower()):
                eliminados.append(d)
                continue
        nuevos.append(d)

    guardar_datos(nuevos)

    if eliminados:
        embed = discord.Embed(
            title="🗑️ Registros eliminados",
            color=discord.Color.red()
        )

        texto = ""
        for e in eliminados:
            texto += f"{e['usuario']} | {e['producto']} | {e['precio']}€\n"

        embed.description = texto[:4000]

        await ctx.send(embed=embed)

        log_channel = bot.get_channel(LOG_BORRADOS_ID)
        if log_channel:
            await log_channel.send(embed=embed)

    else:
        await ctx.send("❌ No encontrado")

# ---------------- PAGO ----------------
class PagoView(discord.ui.View):
    def __init__(self):
        super().__init__(timeout=None)

    @discord.ui.button(label="PayPal", style=discord.ButtonStyle.primary)
    async def paypal(self, interaction: discord.Interaction, button: discord.ui.Button):
        await interaction.response.send_message("💰 PayPal: TU LINK", ephemeral=True)

    @discord.ui.button(label="Bizum", style=discord.ButtonStyle.success)
    async def bizum(self, interaction: discord.Interaction, button: discord.ui.Button):
        await interaction.response.send_message("📱 Bizum: TU NUMERO", ephemeral=True)

    @discord.ui.button(label="Paysafecard", style=discord.ButtonStyle.secondary)
    async def paysafecard(self, interaction: discord.Interaction, button: discord.ui.Button):
        await interaction.response.send_message("💳 Envía el código por privado", ephemeral=True)

@bot.command()
async def pago(ctx):
    embed = discord.Embed(
        title="💳 Métodos de pago",
        description="Selecciona tu método:",
        color=discord.Color.purple()
    )

    await ctx.send(embed=embed, view=PagoView())

# ---------------- START ----------------
if not TOKEN:
    print("❌ ERROR: No hay DISCORD_TOKEN")
    exit()

keep_alive()
bot.run(TOKEN)
